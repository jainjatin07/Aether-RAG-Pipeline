import os
import re
import json
import shutil
import gc
from typing import List
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

# Langchain and MistralAI imports
from langchain_mistralai import MistralAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_mistralai import ChatMistralAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader

# Load environment variables
load_dotenv()

app = FastAPI(title="AetherRAG - Obsidian Knowledge Platform")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directories setup
# Determine base data directory (use 'data' subdirectory in production/Docker, current dir in local dev)
# This allows mounting a single persistent volume to /app/data on Render or Railway.
DATA_ROOT = "data" if os.environ.get("ENV") == "production" else ""

UPLOADS_DIR = os.path.join(DATA_ROOT, "uploads") if DATA_ROOT else "uploads"
CHROMA_DIR = os.path.join(DATA_ROOT, "chroma_db") if DATA_ROOT else "chroma_db"

METADATA_FILE = os.path.join(UPLOADS_DIR, "metadata.json")
os.makedirs(UPLOADS_DIR, exist_ok=True)

# Initialize Metadata JSON if missing
if not os.path.exists(METADATA_FILE):
    with open(METADATA_FILE, "w", encoding="utf-8") as f:
        json.dump([], f)

# Restore Chroma DB from backup if the persistent directory is empty (common on persistent volumes)
BACKUP_DIR = "chroma_db_backup"
if not os.path.exists(CHROMA_DIR) or not os.listdir(CHROMA_DIR):
    if os.path.exists(BACKUP_DIR) and os.listdir(BACKUP_DIR):
        print("Restoring database from backup directory...")
        os.makedirs(CHROMA_DIR, exist_ok=True)
        for item in os.listdir(BACKUP_DIR):
            s = os.path.join(BACKUP_DIR, item)
            d = os.path.join(CHROMA_DIR, item)
            if os.path.isdir(s):
                shutil.copytree(s, d, dirs_exist_ok=True)
            else:
                shutil.copy2(s, d)

# Initialize models
embedding_model = MistralAIEmbeddings(model="mistral-embed")
llm = ChatMistralAI(model="mistral-small-2506")

# Helper: Get on-demand Chroma client to minimize RAM footprint
def get_chroma_vectorstore(collection_name: str) -> Chroma:
    return Chroma(
        persist_directory=CHROMA_DIR,
        embedding_function=embedding_model,
        collection_name=collection_name
    )

# Text Splitter for uploads
splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=250,
)

# Prompt template
prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            """You are a helpful AI assistant.

Use the provided context to answer the question as completely as possible. 

If any part of the question cannot be answered using the provided context, answer the parts that are present, and politely state which details were missing from the documents.
"""
        ),
        (
            "human",
            """Context:
{context}

Question:
{question}
"""
        )
    ]
)

# Helper: Sanitize filename
def secure_filename(filename: str) -> str:
    filename = os.path.basename(filename)
    filename = re.sub(r'[^a-zA-Z0-9._-]', '_', filename)
    return filename

# Helper: Load uploaded files depending on extension
def load_uploaded_file(file_path: str, filename: str) -> List[Document]:
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = filename
        return docs
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            return [Document(page_content=text, metadata={"source": filename})]
        except ImportError:
            raise HTTPException(status_code=500, detail="docx support not configured. Missing python-docx package.")
    else:
        # Default fallback for TXT, MD, PY, JS, JSON etc.
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return [Document(page_content=text, metadata={"source": filename})]
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to read text file: {str(e)}")

# Helper: Load metadata dictionary
def get_metadata() -> dict:
    try:
        with open(METADATA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, list):
                # Migrated old list format to dict format
                return {filename: "ready" for filename in data}
            return data
    except Exception:
        return {}

# Helper: Save metadata dictionary
def save_metadata(metadata: dict):
    with open(METADATA_FILE, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=4)

# Helper: Load active list of successfully indexed files
def get_uploaded_filenames() -> List[str]:
    metadata = get_metadata()
    return [filename for filename, status in metadata.items() if status == "ready"]

# Helper: Save active list of files (backward compatibility)
def save_uploaded_filenames(filenames: List[str]):
    metadata = get_metadata()
    # Keep only the ones in filenames, set to "ready"
    updated_metadata = {fname: metadata.get(fname, "ready") for fname in filenames}
    save_metadata(updated_metadata)

# Helper: Rebuild user Chroma collection from disk files
def rebuild_user_collection(remaining_files: List[str]):
    try:
        db = get_chroma_vectorstore("user_uploads")
        db.delete_collection()
        del db
        gc.collect()
    except Exception:
        pass
    
    db = get_chroma_vectorstore("user_uploads")
    for fname in remaining_files:
        fpath = os.path.join(UPLOADS_DIR, fname)
        if os.path.exists(fpath):
            try:
                docs = load_uploaded_file(fpath, fname)
                if docs:
                    chunks = splitter.split_documents(docs)
                    db.add_documents(chunks)
                    del chunks
                    del docs
                    gc.collect()
            except Exception as e:
                print(f"Error indexing {fname}: {e}")
    del db
    gc.collect()

# API: Serve SPA root
@app.get("/", response_class=HTMLResponse)
def get_dashboard():
    index_path = os.path.join("static", "index.html")
    if not os.path.exists(index_path):
        raise HTTPException(status_code=404, detail="Frontend static index.html not found.")
    with open(index_path, "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

# Helper: Background task to parse and index document
def index_document_task(file_path: str, filename: str):
    try:
        docs = load_uploaded_file(file_path, filename)
        if not docs:
            raise ValueError("Uploaded file is empty.")
        
        # Split documents
        chunks = splitter.split_documents(docs)
        
        # Add to Chroma
        db = get_chroma_vectorstore("user_uploads")
        db.add_documents(chunks)
        
        # Clean up database client instance and garbage collect immediately
        del db
        del chunks
        del docs
        gc.collect()
        
        # Update metadata to mark as ready
        metadata = get_metadata()
        metadata[filename] = "ready"
        save_metadata(metadata)
        print(f"Successfully indexed {filename} in background.")
    except Exception as e:
        print(f"Error indexing {filename} in background: {e}")
        # Remove from metadata on failure
        metadata = get_metadata()
        if filename in metadata:
            del metadata[filename]
            save_metadata(metadata)
        # Cleanup file from disk
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception:
                pass
        gc.collect()

# API: Upload file
@app.post("/api/upload")
async def upload_file(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    filename = secure_filename(file.filename)
    if not filename:
        raise HTTPException(status_code=400, detail="Invalid filename.")
    
    file_path = os.path.join(UPLOADS_DIR, filename)
    
    # Save file on disk
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
    
    # Add to metadata with "indexing" status
    metadata = get_metadata()
    metadata[filename] = "indexing"
    save_metadata(metadata)
    
    # Queue background task for parsing and embedding
    background_tasks.add_task(index_document_task, file_path, filename)
    
    return {"filename": filename, "status": "indexing"}

# API: List uploaded files
@app.get("/api/documents")
def list_documents():
    metadata = get_metadata()
    return [{"filename": k, "status": v} for k, v in metadata.items()]

# API: Delete single document
@app.delete("/api/documents/{filename}")
def delete_document(filename: str):
    filename = secure_filename(filename)
    file_path = os.path.join(UPLOADS_DIR, filename)
    
    metadata = get_metadata()
    if filename not in metadata:
        raise HTTPException(status_code=404, detail="Document not found.")
    
    # Remove file from disk
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to remove file from disk: {str(e)}")
            
    # Remove file reference from metadata
    del metadata[filename]
    save_metadata(metadata)
    
    # Try deleting from collection using metadata filter
    try:
        db = get_chroma_vectorstore("user_uploads")
        db.delete(where={"source": filename})
        del db
    except Exception:
        # Fallback: recreate the collection if delete filter fails
        rebuild_user_collection(get_uploaded_filenames())
        
    gc.collect()
    return {"message": f"Successfully deleted {filename}"}

# API: Clear all uploaded documents
@app.post("/api/reset")
def reset_database():
    metadata = get_metadata()
    files = list(metadata.keys())
    
    # Delete uploaded files from disk
    for fname in files:
        fpath = os.path.join(UPLOADS_DIR, fname)
        if os.path.exists(fpath):
            try:
                os.remove(fpath)
            except Exception:
                pass
                
    # Clear metadata
    save_metadata({})
    
    # Clear collection
    try:
        db = get_chroma_vectorstore("user_uploads")
        db.delete_collection()
        del db
    except Exception as e:
        pass
        
    gc.collect()
    return {"message": "Workspace documents and vector store reset completed."}

# Helper: Split complex multi-part queries into simple search queries
def generate_sub_queries(query: str) -> List[str]:
    prompt_text = f"""Deconstruct the following user query into 1 to 3 simple search queries for a vector database.
Each sub-query must focus on a single concept, topic, question, or term.
Return ONLY a valid JSON list of strings. Do not include markdown code block syntax (like ```json), explanations, or notes.

User Query: "{query}"

Output:"""
    try:
        response = llm.invoke(prompt_text)
        text = response.content.strip()
        # Strip potential markdown wrapping
        if text.startswith("```"):
            lines = text.split("\n")
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines[-1].strip() == "```":
                lines = lines[:-1]
            text = "\n".join(lines).strip()
        
        queries = json.loads(text)
        if isinstance(queries, list) and all(isinstance(q, str) for q in queries):
            return queries
    except Exception as e:
        print(f"Error generating sub-queries: {e}")
    return [query]

# Request schema for querying
class QueryRequest(BaseModel):
    query: str
    mode: str

# API: Query active collections
@app.post("/api/query")
def query_rag(request: QueryRequest):
    query = request.query
    mode = request.mode
    
    # Deconstruct query into sub-queries for multi-document retrieval
    sub_queries = generate_sub_queries(query)
    if query not in sub_queries:
        sub_queries.append(query)
        
    retrieved_docs = []
    
    # Initialize DBs on demand
    default_db = get_chroma_vectorstore("langchain")
    user_db = get_chroma_vectorstore("user_uploads")
    
    # Set up retrievers (lower k per query since we run multiple sub-queries)
    default_retriever = default_db.as_retriever(
        search_type="mmr",
        search_kwargs={"k": 3, "fetch_k": 8, "lambda_mult": 0.5}
    )
    user_retriever = user_db.as_retriever(
        search_type="mmr",
        search_kwargs={"k": 4, "fetch_k": 10, "lambda_mult": 0.5}
    )
    
    # Check collections counts
    try:
        default_count = default_db._collection.count()
    except Exception:
        default_count = 0
        
    try:
        user_count = user_db._collection.count()
    except Exception:
        user_count = 0
        
    # Query database for each sub-query and pool unique chunks
    seen_contents = set()
    
    for sq in sub_queries:
        if mode == "default" or mode == "combined":
            if default_count > 0:
                for doc in default_retriever.invoke(sq):
                    if doc.page_content not in seen_contents:
                        seen_contents.add(doc.page_content)
                        retrieved_docs.append(doc)
                        
        if mode == "user" or mode == "combined":
            if user_count > 0:
                for doc in user_retriever.invoke(sq):
                    if doc.page_content not in seen_contents:
                        seen_contents.add(doc.page_content)
                        retrieved_docs.append(doc)
                        
    # Check if we have anything
    if not retrieved_docs:
        del default_db, user_db, default_retriever, user_retriever
        gc.collect()
        if mode == "default" and default_count == 0:
            raise HTTPException(status_code=400, detail="Default book database is empty or not created.")
        elif mode == "user" and user_count == 0:
            return {"response": "I could not find any uploaded documents in your workspace. Please upload some files (PDF, DOCX, TXT) in the left sidebar first!"}
        return {"response": "No relevant documents could be found for your query. Please refine your question or add more documents."}
        
    # Compile context and extract page citations
    context_parts = []
    citations = []
    seen_citations = set()
    
    for doc in retrieved_docs:
        raw_source = doc.metadata.get("source", "Unknown Source")
        filename = os.path.basename(raw_source)
        source_name = doc.metadata.get("title") or os.path.splitext(filename)[0]
        
        page_val = doc.metadata.get("page")
        if page_val is not None:
            try:
                page_num = int(float(page_val)) + 1
            except (ValueError, TypeError):
                page_num = page_val
        else:
            page_num = None
            
        citation_key = (source_name, page_num)
        if citation_key not in seen_citations:
            seen_citations.add(citation_key)
            citations.append({
                "source": source_name,
                "page": page_num
            })
            
        page_str = f"\nPage {page_num}" if page_num is not None else ""
        context_parts.append(f"Source:\n{source_name}{page_str}\nContent:\n{doc.page_content}")
        
    context = "\n\n".join(context_parts)
    
    # Clean up DB client instances immediately after retrieval is done (before LLM call)
    del default_db, user_db, default_retriever, user_retriever
    gc.collect()
    
    # Run through Mistral AI LLM
    try:
        final_prompt = prompt.invoke({
            "context": context,
            "question": query
        })
        response = llm.invoke(final_prompt)
        
        # Append citations to response
        response_content = response.content.strip()
        if citations:
            citation_blocks = []
            for cit in citations:
                source_name = cit["source"]
                page_num = cit["page"]
                cit_str = f"Source:\n{source_name}"
                if page_num is not None:
                    cit_str += f"\nPage {page_num}"
                citation_blocks.append(cit_str)
            response_content += "\n\n" + "\n\n".join(citation_blocks)
            
        return {"response": response_content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Mistral AI query execution failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    # Dynamic port and host for cloud deployment
    port = int(os.environ.get("PORT", 7860))
    host = os.environ.get("HOST", "127.0.0.1")
    reload = os.environ.get("ENV", "development") == "development"
    uvicorn.run("main:app", host=host, port=port, reload=reload)